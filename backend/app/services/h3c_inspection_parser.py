# -*- coding: utf-8 -*-
"""
H3C 设备巡检解析引擎（平台化移植版）
====================================
功能：解析 H3C 设备 show 命令采集输出，生成巡检 Excel 汇总表和 Word 巡检报告。

本模块从原离线工具 inspection.py 移植而来，去除了 CLI/GUI 入口、多进程调度、
Windows 硬编码路径和 win32com 依赖，供 AIOps 后端服务调用。

依赖：pandas, openpyxl, python-docx
"""

import os
import re
import datetime
from collections import OrderedDict

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION


# ============================================================================
# 模块级预编译正则表达式（优化项1：全局缓存，避免每次解析重复编译）
# ============================================================================

# split_sections 用的提示符模式
_RE_PROMPT = re.compile(r'^<(\S+)>(.+)', re.MULTILINE)
# 兼容两种采集回显：
#   1) 带提示符前缀：  <hostname>display cpu
#   2) 裸命令行：       display cpu   （SSH 回显未带提示符时）
# 两种情况都能正确切分命令段落。
_RE_CMD_START = re.compile(
    r'^(?:<[^>]+>)?'
    r'(?:display|screen-length|quit)\b',
    re.MULTILINE | re.IGNORECASE,
)
_RE_HOSTNAME = re.compile(r'^<([^>]+)>', re.MULTILINE)

# parse_version（Release 前禁止跨行，避免 "Version 5.20\nRelease Version:..." 误匹配）
_RE_VERSION = re.compile(
    r'Version\s+([^\s,]+?),[ \t]*(?:(Release|Alpha|Beta|Build|Version)[ \t]+)?(\S+)',
    re.IGNORECASE,
)
_RE_VERSION_NOCMMA = re.compile(
    r'Version\s+([^\s,]+?)[ \t]+(?:Release|Alpha|Beta|Build)[ \t]+(\S+)',
    re.IGNORECASE,
)
# 从 boot/system 镜像文件名提取版本（-r6328p03.bin / -t7064p15.bin）
_RE_RELEASE_IN_IMAGE = re.compile(r'-([rt]\d+[A-Za-z0-9]*)\.bin', re.IGNORECASE)
_RE_H3C_MODEL = re.compile(r'(H3C\s+\S+)\s+uptime', re.IGNORECASE)
_RE_RELEASE_MODEL = re.compile(
    r'Release\s+Version:\s+H3C\s+(\S+?)(?:-\d+P\d+)?$', re.MULTILINE
)
_RE_MODEL_CLEAN = re.compile(r'-\d{3,}.*$')
_RE_BOARD_TYPE = re.compile(r'BOARD TYPE:\s+(\S+)')

# parse_device — 已改为逐行解析，无需预编译

# parse_cpu
_RE_CPU_CHASSIS = re.compile(
    r'Chassis\s+(\S+)\s+Slot\s+(\S+)\s+CPU\s+(\S+)\s+CPU usage:\s*\n'
    r'\s*(\S+)%\s+in\s+last\s+5\s+seconds\s*\n'
    r'\s*(\S+)%\s+in\s+last\s+1\s+minute\s*\n'
    r'\s*(\S+)%\s+in\s+last\s+5\s+minutes',
    re.MULTILINE
)
_RE_CPU_SLOT = re.compile(
    r'^Slot\s+(\S+)\s+CPU\s+(\S+)\s+CPU usage:\s*\n'
    r'\s*(\S+)%\s+in\s+last\s+5\s+seconds\s*\n'
    r'\s*(\S+)%\s+in\s+last\s+1\s+minute\s*\n'
    r'\s*(\S+)%\s+in\s+last\s+5\s+minutes',
    re.MULTILINE
)
_RE_CPU_UNIT = re.compile(
    r'Unit\s+CPU\s+usage:\s*\n'
    r'\s*(\S+)%\s+in\s+last\s+5\s+seconds\s*\n'
    r'\s*(\S+)%\s+in\s+last\s+1\s+minute\s*\n'
    r'\s*(\S+)%\s+in\s+last\s+5\s+minutes',
    re.MULTILINE
)
# 单行格式（MSR/SecPath/V5 等）："CPU usage: 12% in last 5 seconds"
_RE_CPU_SINGLE = re.compile(
    r'CPU\s+(?:usage|utilization)[^\n]*?(\S+)%\s+in\s+last\s+5\s+seconds',
    re.MULTILINE | re.IGNORECASE
)

# parse_memory
_RE_MEM_CHASSIS_SLOT = re.compile(
    r'Chassis\s+(\S+)\s+Slot\s+(\S+):\s*\n'
    r'(?:(?!Chassis\s).)*?Mem:[^\n]*?(\S+)%',
    re.MULTILINE | re.DOTALL
)
_RE_MEM_SLOT = re.compile(r'^Slot\s+(\S+):\s*$', re.MULTILINE)
_RE_MEM_LINE = re.compile(r'Mem:\s+\S+\s+\S+\s+\S+\s+\S+\s+\S+\s+\S+\s+(\S+)%')
# 单行格式（MSR/SecPath/V5 等）："Memory utilization: 40%"（已用率）
_RE_MEM_SINGLE = re.compile(
    r'Memory\s+(?:utilization|usage)[^\n]*?(\S+)%',
    re.MULTILINE | re.IGNORECASE
)

# parse_power
_RE_PWR_CHASSIS = re.compile(
    r'Chassis\s+(\S+):\s*\n\s*Input Power:.*?\n'
    r'((?:\s*\d+\s+\S+\s+\S+.*?\n)+)',
    re.MULTILINE
)

# parse_fan
# Pattern A: "Fan N State: Normal" (flat, e.g. WX3540X / AC controllers)
_RE_FAN_FLAT = re.compile(r'^Fan\s+(\d+)\s+State:\s*(\S+)', re.MULTILINE)
# Pattern B: "Fan Frame N State: Normal" (e.g. S7606 / S7503E-M)
_RE_FAN_SIMPLE = re.compile(r'Fan Frame\s+(\S+)\s+State:\s*(\S+)', re.MULTILINE)
# Pattern C: "Chassis N: \n   Fan Frame M State: Normal" (legacy chassis)
_RE_FAN_CHASSIS = re.compile(
    r'Chassis\s+(\S+):\s*\n\s*Fan Frame\s+(\S+)\s+State:\s*(\S+)',
    re.MULTILINE
)

# parse_link_aggregation_verbose
_RE_AGG_BLOCK = re.compile(r'\n(?=Aggregate Interface: )')
_RE_AGG_IFACE = re.compile(r'Aggregate Interface:\s*(\S+)')
_RE_AGG_LOCAL = re.compile(
    r'Local:\s*\n(.*?)(?=\nRemote:|\n\s*\nAggregate|\Z)', re.DOTALL
)
_RE_AGG_PORT_TABLE = re.compile(
    r'^\s*Port\s+Status\s+Priority\s+Oper-Key\s*\n(.*?)(?=\n\s*\n[A-Z]|\n\s*\nAggregate|\Z)',
    re.DOTALL | re.MULTILINE
)
_RE_AGG_PORT_DYNAMIC = re.compile(r'^(\S+)\s+([SU])\s+\d+\s+\d+\s+\d+\s+\{[A-Z]*\}')
_RE_AGG_PORT_STATIC = re.compile(r'^(\S+)\s+([SU])\s+\d+\s+\d+\s*$')

# parse_transceiver_diag — 改为按段落分割，仅保留简单数据行正则
_RE_TRX_BLOCK = re.compile(r'\n(?=\S+ transceiver diagnostic)')
_RE_TRX_DATA = re.compile(r'\d+\.\d+')

# parse_bfd_session — 逐行解析，仅作备用
# parse_mac_move — 逐行解析

# parse_interface_status
_RE_IFACE_STATE = re.compile(r'^(\S+)\s*\n\s*Current state:\s*(\S.*)', re.MULTILINE)

# parse_counters_rate / parse_counters — 逐行解析

# parse_logbuffer（新方案）
_RE_LOG_SEV = re.compile(r'/(\d+)/')

# extract_device_model_from_name
_RE_MODEL_FROM_NAME = re.compile(
    r'(S\d{4}(?:X|-EI)?|75\d{3}E(?:-M)?|MSR\d+-\d+|SR\d+)', re.IGNORECASE
)

# version short-name mapping
_RE_AGG_SHORT_BRIDGE = re.compile(r'^bridge-Aggregation(\d+)$', re.IGNORECASE)
_RE_AGG_SHORT_ROUTE = re.compile(r'^route-Aggregation(\d+)$', re.IGNORECASE)


# ============================================================================
# 第一部分：原始文本解析引擎
# ============================================================================

def split_sections_offsets(content):
    """将原始采集文本按命令提示符拆分为偏移量列表（优化项2：延迟切片）。

    兼容两类回显格式：
      - 带提示符前缀：  <hostname>display cpu
      - 裸命令行：       display cpu
    仅当行首为「可选提示符 + 命令动词」时才作为新段落起点，
    提示符独占行（如 <hostname>）不会误判为命令。

    返回:
        offsets: [(command, start_offset, end_offset), ...]
        hostname: 首次出现的设备主机名
    """
    hostname_match = _RE_HOSTNAME.search(content)
    hostname = hostname_match.group(1) if hostname_match else ''

    matches = list(_RE_CMD_START.finditer(content))
    if not matches:
        return [], hostname

    offsets = []
    for i, m in enumerate(matches):
        line_start = m.start()
        line_end = content.find('\n', line_start)
        if line_end == -1:
            line_end = len(content)
        line = content[line_start:line_end]
        # 去掉可选的 <hostname> 前缀，得到纯净命令
        cmd = re.sub(r'^<[^>]+>', '', line).strip()
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
        offsets.append((cmd, start, end))
    return offsets, hostname


def extract_device_model_from_name(sys_name):
    """从设备系统名中提取设备型号。"""
    m = _RE_MODEL_FROM_NAME.search(sys_name)
    if m:
        return m.group(1).upper()
    num_match = re.search(r'(?:^|[^\w])(76\d{2})(?:$|[^\w])', sys_name)
    if num_match:
        return f"S{num_match.group(1)}"
    parts = sys_name.split('-')
    if len(parts) >= 2:
        return parts[-1] if parts[-1] else parts[-2]
    return sys_name


def parse_version(section_text):
    """解析 display version（强制性从本命令提取软件版本与设备型号）。

    软件版本：
      1) "Version 7.1.045, Release 3116"（含无逗号变体）
      2) 镜像文件名 -r6328p03.bin
    设备型号（按优先级）：
      1) "H3C S5130-52S-EI uptime ..."
      2) "Release Version: H3C S5130-52S-EI-3116"（清理尾部版本号）
      3) "BOARD TYPE: S7506E"
    """
    info = {}
    # ---- H3C / 通用格式 ----
    m1 = _RE_VERSION.search(section_text)
    m2 = _RE_VERSION_NOCMMA.search(section_text) if not m1 else None
    if m1:
        info['系统版本'] = f"V{m1.group(1).split('.')[0]}"
        word = (m1.group(2) or "").lower()
        num = m1.group(3) or m1.group(2) or ""
        if word in ("alpha", "beta", "build", "version"):
            info['软件版本'] = f"{m1.group(2).capitalize()}{num}"
        else:
            info['软件版本'] = 'R' + num
    elif m2:
        info['系统版本'] = f"V{m2.group(1).split('.')[0]}"
        info['软件版本'] = 'R' + m2.group(2)
    if not info.get('软件版本'):
        img_match = _RE_RELEASE_IN_IMAGE.search(section_text)
        if img_match:
            ver_num = img_match.group(1)
            info['软件版本'] = ('T' if ver_num.lower().startswith('t') else 'R') + ver_num[1:]

    model = None
    h3c_model = _RE_H3C_MODEL.search(section_text)
    if h3c_model and 'COMWARE' not in h3c_model.group(1).upper():
        model = h3c_model.group(1)
    if not model:
        rel_match = _RE_RELEASE_MODEL.search(section_text)
        if rel_match:
            model = _RE_MODEL_CLEAN.sub('', rel_match.group(1))
    if not model:
        board_match = _RE_BOARD_TYPE.search(section_text)
        if board_match:
            model = board_match.group(1)
    if model:
        info['设备型号'] = model

    board_match = _RE_BOARD_TYPE.search(section_text)
    if board_match:
        info['板卡类型'] = board_match.group(1)
    return info


def parse_device(section_text):
    """解析 display device（优化项5：过滤 NONE/Absent 空槽位）。"""
    boards = []
    lines = section_text.split('\n')
    in_table = False
    has_subslot = False
    for line in lines:
        line = line.strip()
        if not line or line.startswith('<') or line.startswith('Slot'):
            if 'Slot' in line and 'Type' in line and ('State' in line or 'Status' in line):
                in_table = True
                has_subslot = 'Subslot No.' in line
            continue
        if in_table:
            parts = line.split()
            if len(parts) >= 3:
                slot = parts[0]
                if '/' in slot or slot.isdigit():
                    if has_subslot and len(parts) >= 5:
                        # 格式: Slot# Subslot# BoardType Status MaxPorts
                        board_type = parts[2]
                        state = parts[3]
                    else:
                        board_type = parts[1]
                        state = parts[2]
                    # 过滤 NONE 和 Absent 空槽位
                    if board_type.upper() != 'NONE' and state.upper() != 'ABSENT':
                        boards.append({'Slot': slot, 'Type': board_type, 'State': state})
                else:
                    in_table = False
    return boards


def parse_cpu(section_text):
    """解析 display cpu。"""
    cpus = []
    for m in _RE_CPU_CHASSIS.finditer(section_text):
        cpus.append({
            'Chassis': m.group(1), 'Slot': m.group(2), 'Cpu_id': m.group(3),
            'last_5_sec': m.group(4) + '%', 'last_1_min': m.group(5) + '%',
            'last_5_min': m.group(6) + '%',
        })
    for m in _RE_CPU_SLOT.finditer(section_text):
        if not any(c['Slot'] == m.group(1) and c.get('Chassis', '') != '' for c in cpus):
            cpus.append({
                'Chassis': '', 'Slot': m.group(1), 'Cpu_id': m.group(2),
                'last_5_sec': m.group(3) + '%', 'last_1_min': m.group(4) + '%',
                'last_5_min': m.group(5) + '%',
            })
    for m in _RE_CPU_UNIT.finditer(section_text):
        if not cpus:
            cpus.append({
                'Chassis': '', 'Slot': '1', 'Cpu_id': '0',
                'last_5_sec': m.group(1) + '%', 'last_1_min': m.group(2) + '%',
                'last_5_min': m.group(3) + '%',
            })
    # 单行格式兜底（MSR/SecPath/V5 等设备仅输出 5 秒利用率）
    if not cpus:
        m = _RE_CPU_SINGLE.search(section_text)
        if m:
            val = m.group(1) + '%'
            cpus.append({
                'Chassis': '', 'Slot': '1', 'Cpu_id': '0',
                'last_5_sec': val, 'last_1_min': val, 'last_5_min': val,
            })
    return cpus


def parse_memory(section_text):
    """解析 display memory。"""
    memories = []
    for m in _RE_MEM_CHASSIS_SLOT.finditer(section_text):
        chassis = m.group(1)
        slot = m.group(2)
        free_ratio_str = m.group(3)
        try:
            used_ratio = 100.0 - float(free_ratio_str)
        except ValueError:
            used_ratio = 0.0
        memories.append({
            'Chassis': chassis, 'Slot': slot + ':',
            'Cpu_id': '', 'FreeRatio': free_ratio_str + '%',
            'UsedRatio': f"{used_ratio:.1f}%",
        })
    if not memories:
        for m in _RE_MEM_SLOT.finditer(section_text):
            slot = m.group(1)
            block_start = m.end()
            block_end = min(block_start + 500, len(section_text))
            block = section_text[block_start:block_end]
            mem_match = _RE_MEM_LINE.search(block)
            if mem_match:
                free_ratio_str = mem_match.group(1)
                try:
                    used_ratio = 100.0 - float(free_ratio_str)
                except ValueError:
                    used_ratio = 0.0
                memories.append({
                    'Chassis': '', 'Slot': slot + ':',
                    'Cpu_id': '', 'FreeRatio': free_ratio_str + '%',
                    'UsedRatio': f"{used_ratio:.1f}%",
                })
    # 单行格式兜底（MSR/SecPath/V5 等）："Memory utilization: 40%"（已用率）
    if not memories:
        m = _RE_MEM_SINGLE.search(section_text)
        if m:
            try:
                used = float(m.group(1))
            except ValueError:
                used = 0.0
            used = max(0.0, min(used, 100.0))
            memories.append({
                'Chassis': '', 'Slot': '1:', 'Cpu_id': '',
                'FreeRatio': f"{100.0 - used:.1f}%",
                'UsedRatio': f"{used:.1f}%",
            })
    return memories


def parse_power(section_text):
    """解析 display power。"""
    powers = []
    for m in _RE_PWR_CHASSIS.finditer(section_text):
        chassis = m.group(1)
        pwr_block = m.group(2)
        for line in pwr_block.strip().split('\n'):
            parts = line.split()
            if len(parts) >= 2 and parts[0].isdigit():
                power_id = parts[0]
                state = parts[1]
                mode = parts[2] if len(parts) > 2 else ''
                powers.append({
                    'Chassis/Slot': f'Chassis {chassis}',
                    'PowerID': power_id,
                    'State': state,
                    'Mode': mode,
                })
    if not powers:
        lines = section_text.split('\n')
        in_table = False
        for line in lines:
            line_s = line.strip()
            if 'PowerID' in line_s and 'State' in line_s:
                in_table = True
                continue
            if in_table:
                if not line_s or line_s.startswith('<'):
                    break
                parts = line_s.split()
                if len(parts) >= 2 and parts[0].isdigit():
                    powers.append({
                        'Chassis/Slot': '',
                        'PowerID': parts[0],
                        'State': parts[1],
                        'Mode': parts[2] if len(parts) > 2 else '',
                    })

    # --- E. "Device Info on Slot N: / Device ID. Status / N State" (S6800/S6850 等) ---
    if not powers:
        dev_slot = None
        in_dev_table = False
        for line in section_text.split('\n'):
            ls = line.strip()
            dm = re.match(r'^Device Info on Slot\s+(\S+):', ls)
            if dm:
                dev_slot = dm.group(1)
                in_dev_table = False
                continue
            if 'Device ID' in ls and 'Status' in ls:
                in_dev_table = True
                continue
            if in_dev_table:
                if ls.startswith('<'):
                    in_dev_table = False
                    break
                if not ls:
                    continue
                parts = ls.split()
                if len(parts) >= 2 and parts[0].isdigit():
                    powers.append({
                        'Chassis/Slot': f'Slot {dev_slot}',
                        'PowerID': parts[0],
                        'State': parts[1],
                        'Mode': parts[2] if len(parts) > 2 else '',
                    })
    return powers


def parse_fan(section_text):
    """解析 display fan。支持五种格式：
    A. 平铺式:   Fan 1 State: Normal
    B. Fan Frame: Fan Frame 0  State: Normal
    C. Chassis:   Chassis N:\n  Fan Frame M  State: Normal
    D. Slot/Fan:  Slot N:\n  Fan M:\n  State    : Normal (含 Fan-less)
    E. Device Info: Device Info on Slot N:\n  Device ID.  Status\n  1  Normal
    """
    fans = []

    # --- E. "Device Info on Slot N" 格式 (S6800/S6850 等盒式/框式交换机) ---
    dev_slot = None
    in_dev_table = False
    for line in section_text.split('\n'):
        ls = line.strip()
        dm = re.match(r'^Device Info on Slot\s+(\S+):', ls)
        if dm:
            dev_slot = dm.group(1)
            in_dev_table = False
            continue
        if 'Device ID' in ls and 'Status' in ls:
            in_dev_table = True
            continue
        if in_dev_table:
            if ls.startswith('<'):
                in_dev_table = False
                break
            if not ls:
                continue
            parts = ls.split()
            if len(parts) >= 2 and parts[0].isdigit():
                fans.append({
                    'Chassis/Slot': f'Slot {dev_slot}',
                    'Fan_id': parts[0],
                    'Fan_statu': parts[1],
                })
    if fans:
        return fans

    # --- D. Slot/Fan 格式 (按行状态机，覆盖 42 台 S5130/S6520 等) ---
    current_slot = None
    current_fan = None
    for line in section_text.split('\n'):
        line = line.strip()
        if not line:
            continue

        slot_m = re.match(r'^Slot\s+(\S+):\s*$', line)
        if slot_m:
            current_slot = slot_m.group(1)
            current_fan = None
            continue

        if re.match(r'^Fan-less\s*$', line, re.IGNORECASE):
            if current_slot:
                fans.append({
                    'Chassis/Slot': f'Slot {current_slot}',
                    'Fan_id': '--',
                    'Fan_statu': 'Fan-less',
                })
            continue

        fan_m = re.match(r'^Fan\s+(\S+):\s*$', line)
        if fan_m:
            current_fan = fan_m.group(1)
            continue

        state_m = re.match(r'^State\s*:\s*(\S+)', line)
        if state_m and current_slot is not None and current_fan is not None:
            fans.append({
                'Chassis/Slot': f'Slot {current_slot}',
                'Fan_id': current_fan,
                'Fan_statu': state_m.group(1),
            })
            current_fan = None

    # --- A. 平铺式 "Fan N State: Normal" (WX3540X 等 AC 控制器) ---
    if not fans:
        for m in _RE_FAN_FLAT.finditer(section_text):
            fans.append({
                'Chassis/Slot': '',
                'Fan_id': m.group(1),
                'Fan_statu': m.group(2),
            })

    # --- B. "Fan Frame N State: Normal" (S7606 / S7503E-M) ---
    if not fans:
        for m in _RE_FAN_SIMPLE.finditer(section_text):
            fans.append({
                'Chassis/Slot': '',
                'Fan_id': m.group(1),
                'Fan_statu': m.group(2),
            })

    # --- C. Chassis 格式 (兜底) ---
    if not fans:
        for m in _RE_FAN_CHASSIS.finditer(section_text):
            fans.append({
                'Chassis/Slot': f'Chassis {m.group(1)}',
                'Fan_id': m.group(2),
                'Fan_statu': m.group(3),
            })

    return fans


def parse_link_aggregation(section_text):
    """解析 display link-aggregation summary。"""
    aggs = []
    lines = section_text.split('\n')
    in_table = False
    for line in lines:
        line_stripped = line.strip()
        if 'AGG        AGG   Partner ID' in line_stripped:
            in_table = True
            continue
        if in_table:
            if not line_stripped or line_stripped.startswith('<'):
                break
            parts = line_stripped.split()
            if len(parts) >= 6 and parts[0].upper().startswith(('BAGG', 'RAGG')):
                agg_name = parts[0]
                agg_mode = parts[1]
                try:
                    selected = int(parts[-4])
                except (ValueError, IndexError):
                    selected = 0
                try:
                    unselected = int(parts[-3])
                except (ValueError, IndexError):
                    unselected = 0
                individual_info = parts[-2]
                share_type = parts[-1]
                alarm = ''
                if selected > 0 and unselected > 0:
                    alarm = '端口未全选中'
                if selected > 0 and individual_info.upper() == 'U':
                    if not alarm:
                        alarm = 'Individual端口未选中'
                aggs.append({
                    'AGG_Interface': agg_name,
                    'AGG_Mode': agg_mode,
                    'Selected_port': selected,
                    'Unselected_port': unselected,
                    'Individual_type': individual_info,
                    'Share_type': share_type,
                    'Alarm_notes': alarm if alarm else '--',
                    'Unselected_Ports': '',
                    'Port_Status_List': [],
                })
    return aggs


def parse_link_aggregation_verbose(section_text):
    """解析 display link-aggregation verbose。"""
    result = {}
    agg_blocks = _RE_AGG_BLOCK.split(section_text)
    for block in agg_blocks:
        m = _RE_AGG_IFACE.match(block)
        if not m:
            continue
        agg_name = m.group(1)
        local_match = _RE_AGG_LOCAL.search(block)
        if not local_match:
            local_match = _RE_AGG_PORT_TABLE.search(block)
        if not local_match:
            continue
        local_block = local_match.group(1)
        ports = []
        unselected_ports = []
        for line in local_block.split('\n'):
            line_stripped = line.strip()
            port_match = _RE_AGG_PORT_DYNAMIC.match(line_stripped)
            if not port_match:
                port_match = _RE_AGG_PORT_STATIC.match(line_stripped)
            if port_match:
                port_name = port_match.group(1)
                status = port_match.group(2)
                ports.append({'Port': port_name, 'Status': status})
                if status.upper() == 'U':
                    unselected_ports.append(port_name)
        if ports:
            result[agg_name] = {
                'ports': ports,
                'Unselected_Ports': unselected_ports,
            }
    return result


def _parse_transceiver_thresholds(param_lines):
    """从 param_lines 中解析 Alarm thresholds 部分，返回 (rx_low, rx_high, tx_low, tx_high)。
    解析失败返回 (None, None, None, None)，调用方回退硬编码默认值。"""
    # 定位 "Alarm thresholds:" 行
    thresh_start = None
    for i, ln in enumerate(param_lines):
        if ln.strip().startswith('Alarm thresholds'):
            thresh_start = i
            break
    if thresh_start is None:
        return None, None, None, None

    high_line = None
    low_line = None
    for ln in param_lines[thresh_start + 1:]:
        ln = ln.strip()
        if ln.startswith('High'):
            high_line = ln
        elif ln.startswith('Low'):
            low_line = ln

    if not high_line or not low_line:
        return None, None, None, None

    try:
        high_parts = high_line.split()
        low_parts = low_line.split()
        # 列顺序: [High/Low, Temp, Voltage, Bias, RX, TX]
        if len(high_parts) < 6 or len(low_parts) < 6:
            return None, None, None, None
        rx_high = float(high_parts[4])
        tx_high = float(high_parts[5])
        rx_low = float(low_parts[4])
        tx_low = float(low_parts[5])
        return rx_low, rx_high, tx_low, tx_high
    except (ValueError, IndexError):
        return None, None, None, None


def parse_transceiver_diag(section_text):
    """解析 display transceiver diagnosis interface（优化项4：按接口段落分割逐块解析）。"""
    transceivers = []
    blocks = _RE_TRX_BLOCK.split(section_text)
    for block in blocks:
        if not block.strip():
            continue
        # 提取接口名（第一行）
        lines_in_block = block.splitlines()
        if not lines_in_block:
            continue
        first_line = lines_in_block[0].strip()
        if not first_line.endswith('transceiver diagnostic information:'):
            continue
        iface = first_line.split()[0]
        # 跳过 Olt/Onu 接口（不解析光功率）
        if 'Olt' in iface or 'Onu' in iface:
            continue
        if iface.startswith('M-') and not any(c.isdigit() for c in iface.replace('M-', '')[:2]):
            continue
        if 'The transceiver does not support' in block:
            continue

        param_pos = block.find('Current diagnostic parameters:')
        if param_pos == -1:
            continue
        param_block = block[param_pos:]
        param_lines = param_block.splitlines()
        # 跳过表头
        data_line = None
        for ln in param_lines[1:]:
            ln = ln.strip()
            if ln and not ln.startswith('Temp') and not ln.startswith('---'):
                if _RE_TRX_DATA.search(ln):
                    data_line = ln
                    break
        if not data_line:
            continue
        parts = data_line.split()
        if len(parts) < 5:
            continue
        temp, voltage, bias, rx_power, tx_power = parts[:5]

        # 告警判断（使用设备上报阈值 + 3dB 接近临界，解析失败回退默认值）
        rx_alarm = '--'
        tx_alarm = '--'
        alarm_parts = []

        # 解析阈值
        rx_low, rx_high, tx_low, tx_high = _parse_transceiver_thresholds(param_lines)
        # 回退硬编码默认值（保持向后兼容）
        if rx_low is None:
            rx_low = -25.0
        if rx_high is None:
            rx_high = 0.0
        if tx_low is None:
            tx_low = -15.0
        if tx_high is None:
            tx_high = 5.0

        rx_is_na = 'N/A' in rx_power
        tx_is_na = 'N/A' in tx_power
        try:
            if not rx_is_na:
                rx_val = float(rx_power)
                # 接口 down 时的默认读数，非真实光功率异常
                if rx_val not in (-36.96, -40.00):
                    if rx_val < rx_low:
                        rx_alarm = rx_power + 'dBm'
                        alarm_parts.append('收光低于临界')
                    elif rx_val > rx_high:
                        rx_alarm = rx_power + 'dBm'
                        alarm_parts.append('收光高于临界')
            if not tx_is_na:
                tx_val = float(tx_power)
                if tx_val not in (-36.96, -40.00):
                    if tx_val < tx_low:
                        tx_alarm = tx_power + 'dBm'
                        alarm_parts.append('发光低于临界')
                    elif tx_val > tx_high:
                        tx_alarm = tx_power + 'dBm'
                        alarm_parts.append('发光高于临界')
        except (ValueError, TypeError):
            pass
        alarm_notes = '; '.join(alarm_parts) if alarm_parts else '--'

        transceivers.append({
            'Interface': iface,
            'Temp(°C)': temp,
            'Voltage(V)': voltage,
            'Bias(mA)': bias,
            'RX power(dBm)': rx_power,
            'TX power(dBm)': tx_power,
            'RX_Alarm': rx_alarm,
            'TX_Alarm': tx_alarm,
            'Alarm_notes': alarm_notes,
        })
    return transceivers


def parse_bfd_session(section_text):
    """解析 display bfd session。"""
    bfds = []
    lines = section_text.split('\n')
    in_table = False
    for line in lines:
        line_s = line.strip()
        if 'Session State' in line_s and 'Interface' in line_s:
            in_table = True
            continue
        if in_table:
            if not line_s or line_s.startswith('<'):
                break
            parts = line_s.split()
            if len(parts) >= 5:
                state = parts[3] if len(parts) > 3 else ''
                interface = parts[4] if len(parts) > 4 else ''
                alarm = ''
                if state.lower() == 'down':
                    alarm = 'down'
                bfds.append({
                    'BFD_Interface': interface,
                    'State': state,
                    'Alarm_notes': alarm if alarm else '--',
                })
    return bfds


def parse_mac_move(section_text):
    """解析 display mac-address mac-move。"""
    macs = []
    lines = section_text.split('\n')
    in_table = False
    for line in lines:
        line_s = line.strip()
        if 'MAC address' in line_s and 'Current port' in line_s:
            in_table = True
            continue
        if in_table:
            if not line_s or line_s.startswith('<'):
                break
            parts = line_s.split()
            if len(parts) >= 6:
                mac_addr = parts[0]
                vlan = parts[1]
                current_port = parts[2]
                source_port = parts[3]
                move_time = ' '.join(parts[4:6]) if len(parts) >= 6 else ''
                times = parts[-1]
                try:
                    times_int = int(times)
                except ValueError:
                    times_int = 1
                macs.append({
                    'Move_Mac': mac_addr,
                    'Move_Vlan': vlan,
                    'Current_Port': current_port,
                    'Source_Port': source_port,
                    'Move_Time': move_time,
                    'Move_Count': times_int,
                    'Alarm_notes': '漂移',
                })
    return macs


def parse_counters_rate(section_text, direction='inbound'):
    """解析 display counters rate inbound/outbound interface。"""
    interfaces = []
    lines = section_text.split('\n')
    in_table = False
    for line in lines:
        line_s = line.strip()
        if line_s.startswith('Interface') and 'Usage' in line_s:
            in_table = True
            continue
        if in_table:
            if not line_s or line_s.startswith('<') or 'Overflow' in line_s:
                break
            parts = line_s.split()
            if len(parts) >= 2:
                iface = parts[0]
                usage = parts[1]
                if usage.replace('%', '').replace('.', '').isdigit():
                    # 确保末尾有 %
                    if not usage.endswith('%'):
                        usage += '%'
                    interfaces.append({
                        'Interface': iface,
                        'Usage': usage,
                    })
    return interfaces


def parse_counters(section_text):
    """解析 display counters inbound/outbound interface。"""
    interfaces = []
    lines = section_text.split('\n')
    in_table = False
    for line in lines:
        line_s = line.strip()
        if line_s.startswith('Interface') and 'Total' in line_s:
            in_table = True
            continue
        if in_table:
            if not line_s or line_s.startswith('<'):
                break
            parts = line_s.split()
            if len(parts) >= 5:
                iface = parts[0]
                total = parts[1]
                err = parts[-1] if len(parts) >= 5 else '0'
                try:
                    err_int = int(err)
                except ValueError:
                    err_int = 0
                interfaces.append({
                    'Interface': iface,
                    'Total_pkts': total,
                    'Err_pkts': err_int,
                })
    return interfaces


def parse_counters_delta(section_texts):
    """增量计算接口错包数：第二次采集 - 第一次采集。"""
    if not section_texts:
        return []
    if len(section_texts) < 2:
        return parse_counters(section_texts[0])
    first = parse_counters(section_texts[0])
    last = parse_counters(section_texts[-1])
    first_map = {item['Interface']: item['Err_pkts'] for item in first}
    result = []
    for item in last:
        iface = item['Interface']
        err_first = first_map.get(iface, 0)
        err_last = item['Err_pkts']
        delta = err_last - err_first
        result.append({
            'Interface': iface,
            'Total_pkts': item['Total_pkts'],
            'Err_pkts': delta,
        })
    return result


def parse_interface_status(section_text):
    """解析 display interface（P0优化：跳过 Onu 接口）。"""
    interfaces = []
    for m in _RE_IFACE_STATE.finditer(section_text):
        if 'Onu' in m.group(1):
            continue
        interfaces.append({
            'Interface': m.group(1),
            'Statu': m.group(2).upper() if m.group(2).lower() == 'up' else m.group(2).upper(),
        })
    return interfaces


def parse_environment(section_text):
    """解析 display environment。"""
    envs = []
    lines = section_text.split('\n')
    in_table = False
    for line in lines:
        line_s = line.strip()
        if 'Slot' in line_s and 'Sensor' in line_s and 'Temperature' in line_s:
            in_table = True
            continue
        if in_table:
            if line_s.startswith('<'):
                break
            if not line_s:
                continue
            parts = line_s.split()
            if len(parts) >= 6:
                chassis = parts[0]
                slot = parts[1]
                try:
                    temp = float(parts[3])
                    warning = float(parts[5]) if len(parts) > 5 else 100
                    alarm = float(parts[6]) if len(parts) > 6 else 100
                except (ValueError, IndexError):
                    continue
                alarm_note = '--'
                if temp >= alarm:
                    alarm_note = '温度过高告警'
                elif temp >= warning:
                    alarm_note = '温度偏高警告'
                envs.append({
                    'Chassis': chassis,
                    'Slot': slot,
                    'Sensor': parts[2],
                    'Temperature': temp,
                    'Alarm_notes': alarm_note,
                })
    return envs


def parse_device_manuinfo(section_text):
    """解析 display device manuinfo。

    支持三种格式：
    1. 固定机箱（如 WX2580X）：顶层 DEVICE_SERIAL_NUMBER:xxx
    2. 机框设备（如 S7506X/S7506X-G）：Chassis self + Slot X CPU 0 块
    3. 过滤 Fan/Power 块，跳过 DEVICE_SERIAL_NUMBER 为空的条目
    """
    barcodes = []
    lines = section_text.split('\n')

    # 当前块上下文
    current_block = None       # 'Chassis' / 'Slot' / 'Fan' / 'Power' / None(固定机箱)
    current_slot = None        # 槽位号（仅 Slot 块）
    current_name = None
    current_sn = None

    def flush_block():
        """提交当前块到结果列表（含过滤逻辑）。"""
        nonlocal current_block, current_slot, current_name, current_sn
        if current_block in ('Fan', 'Power'):
            # 跳过 Fan/Power 块
            current_block = current_slot = current_name = current_sn = None
            return
        if current_sn:
            slot_label = current_slot if current_block == 'Slot' else ('Chassis' if current_block == 'Chassis' else '--')
            barcodes.append({
                'Slot': slot_label,
                'Device_Name': current_name or '',
                'Info': current_sn,
            })
        current_block = current_slot = current_name = current_sn = None

    for line in lines:
        line_s = line.strip()
        if not line_s:
            continue

        # 检测块头：Chassis self: / Slot X CPU 0: / Fan X: / Power X:
        header_match = re.match(
            r'(Chassis self|Slot\s+\d+\s+CPU\s+\d+|Fan\s+\d+|Power\s+\d+)\s*:\s*$',
            line_s
        )
        if header_match:
            flush_block()
            header = header_match.group(1)
            if header == 'Chassis self':
                current_block = 'Chassis'
            elif header.startswith('Slot'):
                current_block = 'Slot'
                # 提取槽位号，如 "Slot 3 CPU 0" → "3"
                slot_m = re.search(r'Slot\s+(\d+)', header)
                current_slot = slot_m.group(1) if slot_m else ''
            elif header.startswith('Fan'):
                current_block = 'Fan'
            elif header.startswith('Power'):
                current_block = 'Power'
            continue

        # 检测结束条件：<hostname> 提示符
        if line_s.startswith('<'):
            flush_block()
            continue

        # 块内键值对解析
        # 格式1（固定机箱）: DEVICE_SERIAL_NUMBER:210235A...
        # 格式2（机框）    : DEVICE_SERIAL_NUMBER : 210235A...
        kv_match = re.match(r'(DEVICE_NAME|DEVICE_SERIAL_NUMBER)\s*:+\s*(.*)', line_s)
        if not kv_match:
            continue

        key = kv_match.group(1)
        val = kv_match.group(2).strip()
        if key == 'DEVICE_NAME':
            current_name = val
        elif key == 'DEVICE_SERIAL_NUMBER' and val:
            current_sn = val

    # 处理最后一块（未以 < 结束）
    flush_block()
    return barcodes


def parse_logbuffer(section_text, max_lines=500):
    """解析 display logbuffer reverse（优化项3：从尾部向上逐行扫描，最多取 max_lines 条）。

    仅保留 ERROR(3) 及以上级别日志，忽略 WARNING/INFO/DEBUG。
    H3C 日志格式：%Mon DD HH:MM:SS:mmm YYYY Hostname Module/Severity/Submodule: Message
    """
    logs = []
    sev_map = {'0': 'EMERGENCY', '1': 'ALERT', '2': 'CRITICAL', '3': 'ERROR'}
    lines = section_text.splitlines()
    count = 0
    for line in reversed(lines):
        if count >= max_lines:
            break
        line = line.strip()
        if not line.startswith('%'):
            continue
        try:
            parts = line.split(None, 8)
            if len(parts) < 9:
                continue
            # parts[0]='%Jul', parts[1]='9', parts[2]='14:55:12:300',
            # parts[3]='2026', parts[4]='HeXin-S7606-IRF-1',
            # parts[5]='SHELL/6/SHELL_CMD:', parts[6:] message
            month = parts[0][1:]   # strip leading %
            day = parts[1]
            time_str = parts[2]
            year = parts[3]
            hostname = parts[4]
            sev_part = parts[5]
            sev_match = _RE_LOG_SEV.search(sev_part)
            if not sev_match:
                continue
            sev_num = sev_match.group(1)
            if sev_num not in ('0', '1', '2', '3'):
                continue
            severity = sev_map.get(sev_num, 'UNKNOWN')
            timestamp = f"{month} {day} {year} {time_str}"
            message = ' '.join(parts[6:])
            full_message = f"{hostname} {sev_part} {message}"
            logs.append({
                'Severity': severity,
                'Timestamp': timestamp,
                'Message': full_message,
            })
            count += 1
        except Exception:
            continue
    logs.reverse()
    return logs



def detect_smart_anomalies(logbuffer_raw):
    """智能异常检测：从 raw logbuffer 文本中检测接口频繁 up/down 等异常。

    返回: list[dict] 异常摘要列表，
         每个 dict 含 type / interface / summary / details 字段。
    """
    import datetime
    anomalies = []
    if not logbuffer_raw:
        return anomalies

    # ── 接口频繁 up/down 检测 ──
    # H3C 格式: %Mon DD HH:MM:SS:mmm YYYY Hostname IFNET/... interface XXX changed to (up|down)
    updown_pattern = re.compile(
        r'%(\w{3})\s+(\d+)\s+(\d{2}:\d{2}:\d{2}):\d+\s+(\d{4})\s+\S+\s+IFNET/\d+/(?:LINK_UPDOWN|PHY_UPDOWN):\s+.*?interface\s+(\S+)\s+changed to\s+(\w+)',
        re.IGNORECASE
    )

    events = []
    for m in updown_pattern.finditer(logbuffer_raw):
        iface = m.group(5)
        # 忽略 ONU 接口
        if re.search(r'[Oo][Nn][Uu]', iface):
            continue
        month = m.group(1)
        day = m.group(2)
        time_str = m.group(3)
        year = m.group(4)
        state = m.group(6).lower()
        events.append({
            'iface': iface,
            'timestamp': f"{month} {day} {year} {time_str}",
            'state': state,
        })

    # 按接口分组，检测 5 分钟内 >= 3 次 up/down 的簇
    iface_events = {}
    for evt in events:
        iface_events.setdefault(evt['iface'], []).append(evt)

    def _time_to_minutes(ts):
        """''Mon DD YYYY HH:MM:SS'' → 近似分钟数（同天比较）"""
        import datetime
        try:
            dt = datetime.datetime.strptime(ts, '%b %d %Y %H:%M:%S')
            return dt
        except ValueError:
            return None

    for iface, evts in iface_events.items():
        evts.sort(key=lambda e: _time_to_minutes(e['timestamp']) or datetime.datetime(2000, 1, 1))
        reported = False
        for i in range(len(evts)):
            if reported:
                break
            cluster = [evts[i]]
            base_t = _time_to_minutes(evts[i]['timestamp'])
            if base_t is None:
                continue
            for j in range(i + 1, len(evts)):
                t = _time_to_minutes(evts[j]['timestamp'])
                if t is None:
                    continue
                if (t - base_t).total_seconds() <= 300:  # 5 分钟
                    cluster.append(evts[j])
            if len(cluster) >= 3:
                # 过滤超过3个月的异常日志（2026-07-29 前 3 个月 = 2026-04-29）
                if base_t < datetime.datetime.now() - datetime.timedelta(days=90):
                    continue
                start_ts = cluster[0]['timestamp']
                end_ts = cluster[-1]['timestamp']
                anomalies.append({
                    'type': '接口频繁up/down',
                    'interface': iface,
                    'summary': f"{iface} 在 {start_ts} 至 {end_ts} 期间频繁 up/down（{len(cluster)}次）",
                })
                reported = True

    return anomalies
def _read_file_auto_encoding(filepath):
    """自动检测编码读取文件内容。

    策略：UTF-8 优先（严格解码自检），失败则回退 GBK+replace。
    - UTF-8 文件（如部分 CYW 采集文件）：UTF-8 严格解码一次成功。
    - GBK 文件（如 GAW H3C 采集文件）：UTF-8 严格解码失败，
      回退 GBK+replace 容忍局部坏字节，确保 sysname 正确提取。
    - gb18030 虽然覆盖面更广但会误将 UTF-8 字节流解释为"合法"
      gb18030 序列，导致中文乱码，因此不列入候选编码。
    """
    with open(filepath, 'rb') as f:
        raw = f.read()

    # 1. 优先尝试 UTF-8（严格模式：合法 UTF-8 文件一次通过）
    try:
        return raw.decode('utf-8')
    except UnicodeDecodeError:
        pass

    # 2. UTF-8 失败 → GBK 系（errors='replace' 容忍局部坏字节）
    for encoding in ['gbk', 'gb2312']:
        try:
            return raw.decode(encoding, errors='replace')
        except LookupError:
            continue

    # 3. 最终兜底
    return raw.decode('utf-8', errors='replace')



def _extract_ip_from_filename(filename):
    """Extract IP from filename.

    Supports two formats:
    - Old: ``10.10.1.30.txt`` → ``10.10.1.30``
    - New: ``2025-12-18_15.37_10.15.105.10.log`` → ``10.15.105.10``

    Strategy: strip extension first, then walk from rightmost ``_`` segment;
    the first segment that looks like a valid dotted-quad IP is returned.
    If none found, falls back to the full stem.
    """
    ip_pattern = re.compile(r'^(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})$')
    stem = os.path.splitext(filename)[0]
    parts = stem.split('_')
    for part in reversed(parts):
        if ip_pattern.match(part):
            return part
    return stem  # fallback: whole stem (legacy behaviour)
def parse_raw_file(filepath, content=None):
    """解析单个原始采集文件（优化版：基于偏移量延迟切片）。

    Args:
        filepath: 文件路径（用于提取设备标识名）
        content: 可选，预读的文件内容字符串。若未提供则从 filepath 读取。

    Returns:
        结构化的设备数据字典
    """
    if content is None:
        content = _read_file_auto_encoding(filepath)

    cmd_offsets, hostname = split_sections_offsets(content)
    filename = os.path.basename(filepath)

    device_data = {
        'Sys_ip': _extract_ip_from_filename(filename),
        'Sys_name': hostname,
        '系统版本': 'V7',
        '软件版本': '',
        '设备型号': '',
        'cpu_list': [],
        'memory_list': [],
        'board_list': [],
        'power_list': [],
        'fan_list': [],
        'agg_list': [],
        'agg_verbose': {},
        'transceiver_list': [],
        'bfd_list': [],
        'mac_move_list': [],
        'interface_rate_in': [],
        'interface_rate_out': [],
        'interface_counters_in': [],
        'interface_counters_out': [],
        'interface_status_list': [],
        'env_list': [],
        'barcode_list': [],
        'log_list': [],
        'logbuffer_raw': '',
    }

    # 同命令多次采集的段落缓存（为 counters 增量计算保留）
    multi_section_cache = {}

    for cmd, start, end in cmd_offsets:
        cmd_lower = cmd.lower()

        # 跳过 display current-configuration（优化项6）
        if cmd_lower in ('display current-configuration', 'display current-configuration |'):
            continue

        section_text = content[start:end]

        try:
            if 'display version' == cmd_lower:
                ver_info = parse_version(section_text)
                device_data.update(ver_info)
            elif 'display device' == cmd_lower and 'manuinfo' not in cmd_lower:
                device_data['board_list'] = parse_device(section_text)
            elif 'display cpu' == cmd_lower:
                device_data['cpu_list'] = parse_cpu(section_text)
            elif 'display memory' == cmd_lower:
                device_data['memory_list'] = parse_memory(section_text)
            elif 'display power' == cmd_lower and 'verbose' not in cmd_lower:
                device_data['power_list'] = parse_power(section_text)
            elif 'display fan' == cmd_lower:
                device_data['fan_list'] = parse_fan(section_text)
            elif 'link-aggregation summary' in cmd_lower:
                device_data['agg_list'] = parse_link_aggregation(section_text)
            elif 'link-aggregation verbose' in cmd_lower:
                device_data['agg_verbose'] = parse_link_aggregation_verbose(section_text)
            elif 'transceiver diagnosis' in cmd_lower:
                device_data['transceiver_list'] = parse_transceiver_diag(section_text)
            elif 'bfd session' in cmd_lower:
                device_data['bfd_list'] = parse_bfd_session(section_text)
            elif 'mac-address mac-move' in cmd_lower:
                device_data['mac_move_list'] = parse_mac_move(section_text)
            elif 'counters rate inbound' in cmd_lower:
                device_data['interface_rate_in'] = parse_counters_rate(section_text, 'inbound')
            elif 'counters rate outbound' in cmd_lower:
                device_data['interface_rate_out'] = parse_counters_rate(section_text, 'outbound')
            elif 'counters inbound' in cmd_lower and 'rate' not in cmd_lower:
                if cmd_lower not in multi_section_cache:
                    multi_section_cache[cmd_lower] = []
                multi_section_cache[cmd_lower].append(section_text)
            elif 'counters outbound' in cmd_lower and 'rate' not in cmd_lower:
                if cmd_lower not in multi_section_cache:
                    multi_section_cache[cmd_lower] = []
                multi_section_cache[cmd_lower].append(section_text)
            elif 'display interface' == cmd_lower:
                device_data['interface_status_list'] = parse_interface_status(section_text)
            elif 'display environment' == cmd_lower:
                device_data['env_list'] = parse_environment(section_text)
            elif 'device manuinfo' in cmd_lower:
                device_data['barcode_list'] = parse_device_manuinfo(section_text)
            elif 'logbuffer reverse' in cmd_lower:
                device_data['log_list'] = parse_logbuffer(section_text)
                device_data['logbuffer_raw'] = section_text
        except Exception as e:
            print(f"  [WARN] 解析 {cmd} 时出错: {e}")

    # 处理增量 counters（多次采集）
    for cmd_key, texts in multi_section_cache.items():
        try:
            if 'counters inbound' in cmd_key:
                device_data['interface_counters_in'] = parse_counters_delta(texts)
            elif 'counters outbound' in cmd_key:
                device_data['interface_counters_out'] = parse_counters_delta(texts)
        except Exception as e:
            print(f"  [WARN] 解析 {cmd_key} 增量时出错: {e}")

    # 设备型号回退推断
    if not device_data.get('设备型号') or device_data.get('设备型号', '').upper() == 'COMWARE':
        # 优先从单板列表提取（WX 系列等设备需 Board Type 中的完整型号）
        if device_data['board_list']:
            for b in device_data['board_list']:
                if b['Type'] != 'NONE':
                    device_data['设备型号'] = b['Type']
                    break
        # 其次从设备名推断
        if not device_data.get('设备型号'):
            model = extract_device_model_from_name(device_data['Sys_name'])
            device_data['设备型号'] = model

    # 合并 verbose 聚合端口状态到 agg_list
    if device_data['agg_verbose']:
        verbose_map = {}
        for full_name, vdata in device_data['agg_verbose'].items():
            m = _RE_AGG_SHORT_BRIDGE.match(full_name)
            if m:
                short_name = 'BAGG' + m.group(1)
            else:
                m = _RE_AGG_SHORT_ROUTE.match(full_name)
                if m:
                    short_name = 'RAGG' + m.group(1)
                else:
                    short_name = full_name
            verbose_map[short_name] = vdata
        for a in device_data['agg_list']:
            agg_name = a.get('AGG_Interface', '')
            vdata = verbose_map.get(agg_name)
            if vdata:
                a['Unselected_Ports'] = ', '.join(vdata.get('Unselected_Ports', []))
                a['Port_Status_List'] = vdata.get('ports', [])
                if a['Unselected_Ports'] and a['Alarm_notes'] != '--':
                    a['Alarm_notes'] += f" | 未选中端口: {a['Unselected_Ports']}"

    return device_data


# ============================================================================
# 第二部分：汇总数据构建（用于 Excel 生成）
# ============================================================================

def get_max_cpu(cpu_list):
    if not cpu_list:
        return 'N/A', 'N/A'
    max_val = 0
    for c in cpu_list:
        try:
            val = float(c['last_5_min'].replace('%', ''))
            if val > max_val:
                max_val = val
        except (ValueError, KeyError):
            continue
    return f"{max_val:.0f}%", '异常' if max_val > 70 else '正常'


def get_max_memory(memory_list):
    if not memory_list:
        return 'N/A', 'N/A'
    max_val = 0
    for m in memory_list:
        try:
            val = float(m['UsedRatio'].replace('%', ''))
            if val > max_val:
                max_val = val
        except (ValueError, KeyError):
            continue
    return f"{max_val:.1f}%", '异常' if max_val > 80 else '正常'


def check_board_status(board_list):
    if not board_list:
        return '正常'
    for b in board_list:
        state = b['State'].upper()
        btype = b['Type'].upper()
        if state not in ('NORMAL', 'MASTER', 'SLAVE', 'STANDBY', 'ABSENT') and btype != 'NONE':
            return '异常'
    return '正常'


def check_power_status(power_list):
    if not power_list:
        return '正常'
    for p in power_list:
        if p['State'].upper() != 'NORMAL':
            return '异常'
    return '正常'


def check_fan_status(fan_list):
    if not fan_list:
        return '正常'
    for f in fan_list:
        status = f['Fan_statu'].upper()
        if status not in ('NORMAL', 'FAN-LESS'):
            return '异常'
    return '正常'


def check_agg_status(agg_list):
    if not agg_list:
        return '正常'
    for a in agg_list:
        if a['Alarm_notes'] != '--':
            return '异常'
    return '正常'


def check_transceiver_status(transceiver_list):
    if not transceiver_list:
        return '正常'
    for t in transceiver_list:
        if t['Alarm_notes'] != '--' and t['Alarm_notes'] != '需手动查看':
            return '异常'
    return '正常'


def check_interface_usage(rate_list, threshold=60):
    if not rate_list:
        return '正常'
    for r in rate_list:
        try:
            usage = float(r['Usage'].replace('%', ''))
            if usage > threshold:
                return '异常'
        except (ValueError, KeyError):
            continue
    return '正常'


def check_interface_errors(counters_list):
    if not counters_list:
        return '正常'
    for c in counters_list:
        if c['Err_pkts'] > 0:
            return '异常'
    return '正常'


def check_bfd_status(bfd_list):
    if not bfd_list:
        return '无'
    for b in bfd_list:
        if 'down' in b.get('Alarm_notes', '').lower():
            return '异常'
    return '正常'


def check_mac_move(mac_list):
    if not mac_list:
        return '正常'
    return '异常'


def check_logbuffer(log_list):
    if not log_list:
        return '正常'
    if len(log_list) > 0:
        return '异常'
    return '正常'


def build_summary_row(device_data):
    cpu_max, cpu_status = get_max_cpu(device_data['cpu_list'])
    mem_max, mem_status = get_max_memory(device_data['memory_list'])
    return {
        'Sys_ip': device_data['Sys_ip'],
        'Sys_name': device_data['Sys_name'],
        '系统版本': device_data.get('系统版本', 'V7'),
        '软件版本': device_data.get('软件版本', ''),
        '设备型号': device_data.get('设备型号', ''),
        'CPU占用': cpu_status,
        'CPU利用率': cpu_max,
        '内存占用': mem_status,
        '内存利用率': mem_max,
        '设备/板卡状态': check_board_status(device_data['board_list']),
        '电源状态': check_power_status(device_data['power_list']),
        '风扇状态': check_fan_status(device_data['fan_list']),
        '聚合选中状态': check_agg_status(device_data['agg_list']),
        '光功率状态': check_transceiver_status(device_data['transceiver_list']),
        '接口接收利用率': check_interface_usage(device_data['interface_rate_in']),
        '接口发送利用率': check_interface_usage(device_data['interface_rate_out']),
        '接口接收错包': check_interface_errors(device_data['interface_counters_in']),
        '接口发送错包': check_interface_errors(device_data['interface_counters_out']),
        'BFD状态': check_bfd_status(device_data['bfd_list']),
        'MAC漂移': check_mac_move(device_data['mac_move_list']),
        'ERROR及以上日志': check_logbuffer(device_data['log_list']),
    }


# ============================================================================
# 第三部分：Excel 生成
# ============================================================================

THIN_BORDER = Border(
    left=Side(style='thin'),
    right=Side(style='thin'),
    top=Side(style='thin'),
    bottom=Side(style='thin')
)
HEADER_FILL = PatternFill(start_color='ADD8E6', end_color='ADD8E6', fill_type='solid')
HEADER_FONT = Font(name='微软雅黑', size=10, bold=True)
CELL_FONT = Font(name='微软雅黑', size=9)
CENTER_ALIGN = Alignment(horizontal='center', vertical='center', wrap_text=True)


def auto_width(ws, min_width=8, max_width=50):
    for col_idx in range(1, ws.max_column + 1):
        max_len = 0
        col_letter = get_column_letter(col_idx)
        for row_idx in range(1, ws.max_row + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            if cell.value:
                lines = str(cell.value).split('\n')
                for line in lines:
                    max_len = max(max_len, len(line) * 1.2)
        ws.column_dimensions[col_letter].width = min(max(max_len + 2, min_width), max_width)


def write_header(ws, headers, row=1):
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col_idx, value=header)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER


def write_data_rows(ws, data, headers, start_row=2):
    for row_idx, row_data in enumerate(data, start_row):
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=row_data.get(header, ''))
            cell.font = CELL_FONT
            cell.alignment = CENTER_ALIGN
            cell.border = THIN_BORDER


def generate_excel(all_device_data, output_path):
    wb = Workbook()
    wb.remove(wb.active)

    # --- Sheet 1: 汇总信息 ---
    ws_summary = wb.create_sheet('汇总信息')
    summary_headers = [
        'Sys_ip', 'Sys_name', '系统版本', '软件版本', '设备型号',
        'CPU占用', 'CPU利用率', '内存占用', '内存利用率',
        '设备/板卡状态', '电源状态', '风扇状态', '聚合选中状态',
        '光功率状态', '接口接收利用率', '接口发送利用率',
        '接口接收错包', '接口发送错包', 'BFD状态', 'MAC漂移', 'ERROR及以上日志'
    ]
    write_header(ws_summary, summary_headers)
    summary_rows = [build_summary_row(dd) for dd in all_device_data]
    write_data_rows(ws_summary, summary_rows, summary_headers)
    red_font = Font(name='微软雅黑', size=9, color='FF0000', bold=True)
    for row_idx in range(2, ws_summary.max_row + 1):
        for col_idx in range(1, ws_summary.max_column + 1):
            cell = ws_summary.cell(row=row_idx, column=col_idx)
            if cell.value and '异常' in str(cell.value):
                cell.font = red_font
    auto_width(ws_summary)

    # --- Sheet 2: 异常日志信息 ---
    ws_log = wb.create_sheet('异常日志信息')
    log_headers = ['Sys_ip', 'Sys_name', 'Severity', 'Timestamp', 'warning_log']
    write_header(ws_log, log_headers)
    log_row = 2
    for dd in all_device_data:
        for l in dd['log_list']:
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'Severity': l.get('Severity', ''), 'Timestamp': l.get('Timestamp', ''),
                'warning_log': l.get('Message', ''),
            }
            write_data_rows(ws_log, [row_data], log_headers, log_row)
            log_row += 1
    auto_width(ws_log)

    # --- Sheet 2.5: 全部日志信息 ---
    ws_log_all = wb.create_sheet('全部日志信息')
    log_all_headers = ['Sys_ip', 'Sys_name', '日志原文']
    write_header(ws_log_all, log_all_headers)
    log_all_row = 2
    for dd in all_device_data:
        raw_text = dd.get('logbuffer_raw', '')
        if raw_text:
            # 截取前 30000 字符避免 Excel 单元格超限
            raw_text = raw_text[:30000]
        ws_log_all.cell(row=log_all_row, column=1, value=dd['Sys_ip'])
        ws_log_all.cell(row=log_all_row, column=2, value=dd['Sys_name'])
        cell = ws_log_all.cell(row=log_all_row, column=3, value=raw_text)
        cell.alignment = Alignment(wrap_text=True, vertical='top')
        # 设置行高自适应
        ws_log_all.row_dimensions[log_all_row].height = 200
        log_all_row += 1
    auto_width(ws_log_all)

    # --- Sheet 3: 接口信息 ---
    ws_iface = wb.create_sheet('接口信息')
    iface_headers = ['Sys_ip', 'Sys_name', 'Interface', 'In_usage', 'Out_usage',
                     'In_error_count', 'Out_error_count', 'Statu',
                     'In_error_add', 'Out_error_add', 'In_usage_warning', 'Out_usage_warning', 'Alarm_notes']
    write_header(ws_iface, iface_headers)
    iface_row = 2
    for dd in all_device_data:
        in_rate_map = {r['Interface']: r['Usage'] for r in dd['interface_rate_in']}
        out_rate_map = {r['Interface']: r['Usage'] for r in dd['interface_rate_out']}
        in_err_map = {r['Interface']: r['Err_pkts'] for r in dd['interface_counters_in']}
        out_err_map = {r['Interface']: r['Err_pkts'] for r in dd['interface_counters_out']}
        status_map = {r['Interface']: r['Statu'] for r in dd['interface_status_list']}

        all_ifaces = set()
        all_ifaces.update(in_rate_map.keys(), out_rate_map.keys(), in_err_map.keys(), out_err_map.keys(), status_map.keys())

        for iface in sorted(all_ifaces):
            in_usage = in_rate_map.get(iface, 'N/A')
            out_usage = out_rate_map.get(iface, 'N/A')
            in_err = in_err_map.get(iface, 0)
            out_err = out_err_map.get(iface, 0)
            status = status_map.get(iface, 'N/A')
            in_warning = '--'
            out_warning = '--'
            try:
                if float(in_usage.replace('%', '')) > 60:
                    in_warning = in_usage
            except (ValueError, KeyError):
                pass
            try:
                if float(out_usage.replace('%', '')) > 60:
                    out_warning = out_usage
            except (ValueError, KeyError):
                pass
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'], 'Interface': iface,
                'In_usage': in_usage, 'Out_usage': out_usage,
                'In_error_count': in_err, 'Out_error_count': out_err, 'Statu': status,
                'In_error_add': str(in_err) if in_err > 0 else '--',
                'Out_error_add': str(out_err) if out_err > 0 else '--',
                'In_usage_warning': in_warning, 'Out_usage_warning': out_warning,
                'Alarm_notes': '',
            }
            write_data_rows(ws_iface, [row_data], iface_headers, iface_row)
            iface_row += 1
    auto_width(ws_iface)

    # --- Sheet 4: Cpu状态 ---
    ws_cpu = wb.create_sheet('Cpu状态')
    cpu_headers = ['Sys_ip', 'Sys_name', 'Chassis', 'Slot', 'Cpu_id', 'last_5_sec', 'last_1_min', 'last_5_min', 'Alarm_notes']
    write_header(ws_cpu, cpu_headers)
    cpu_row = 2
    for dd in all_device_data:
        for c in dd['cpu_list']:
            alarm = '--'
            try:
                if float(c.get('last_5_min', '0%').replace('%', '')) > 70:
                    alarm = 'CPU占用过高'
            except ValueError:
                pass
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'Chassis': c.get('Chassis', ''), 'Slot': c.get('Slot', ''),
                'Cpu_id': c.get('Cpu_id', ''), 'last_5_sec': c.get('last_5_sec', ''),
                'last_1_min': c.get('last_1_min', ''), 'last_5_min': c.get('last_5_min', ''),
                'Alarm_notes': alarm,
            }
            write_data_rows(ws_cpu, [row_data], cpu_headers, cpu_row)
            cpu_row += 1
    auto_width(ws_cpu)

    # --- Sheet 5: 内存状态 ---
    ws_mem = wb.create_sheet('内存状态')
    mem_headers = ['Sys_ip', 'Sys_name', 'Chassis', 'Slot', 'Cpu_id', 'FreeRatio', 'Alarm_notes']
    write_header(ws_mem, mem_headers)
    mem_row = 2
    for dd in all_device_data:
        for m in dd['memory_list']:
            alarm = '--'
            try:
                if float(m.get('UsedRatio', '0%').replace('%', '')) > 80:
                    alarm = '内存占用过高'
            except ValueError:
                pass
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'Chassis': m.get('Chassis', ''), 'Slot': m.get('Slot', ''),
                'Cpu_id': m.get('Cpu_id', ''), 'FreeRatio': m.get('FreeRatio', ''),
                'Alarm_notes': alarm,
            }
            write_data_rows(ws_mem, [row_data], mem_headers, mem_row)
            mem_row += 1
    auto_width(ws_mem)

    # --- Sheet 6: 单板状态 ---
    ws_board = wb.create_sheet('单板状态')
    board_headers = ['Sys_ip', 'Sys_name', 'Slot', 'Brd_statu', 'Type', 'Alarm_notes']
    write_header(ws_board, board_headers)
    board_row = 2
    for dd in all_device_data:
        for b in dd['board_list']:
            alarm = '--'
            state = b['State'].upper()
            btype = b['Type'].upper()
            if state not in ('NORMAL', 'MASTER', 'SLAVE', 'STANDBY', 'ABSENT') and btype != 'NONE':
                alarm = '单板异常'
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'Slot': b['Slot'], 'Brd_statu': b['State'], 'Type': b['Type'],
                'Alarm_notes': alarm,
            }
            write_data_rows(ws_board, [row_data], board_headers, board_row)
            board_row += 1
    auto_width(ws_board)

    # --- Sheet 7: 电源信息 ---
    ws_pwr = wb.create_sheet('电源信息')
    pwr_headers = ['Sys_ip', 'Sys_name', 'Chassis/Slot', 'PowerID', 'State', 'Mode', 'Alarm_notes']
    write_header(ws_pwr, pwr_headers)
    pwr_row = 2
    for dd in all_device_data:
        for p in dd['power_list']:
            alarm = '--'
            if p['State'].upper() != 'NORMAL':
                alarm = f"电源{p['State']}"
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'Chassis/Slot': p.get('Chassis/Slot', ''), 'PowerID': p.get('PowerID', ''),
                'State': p.get('State', ''), 'Mode': p.get('Mode', ''),
                'Alarm_notes': alarm,
            }
            write_data_rows(ws_pwr, [row_data], pwr_headers, pwr_row)
            pwr_row += 1
    auto_width(ws_pwr)

    # --- Sheet 8: 风扇信息 ---
    ws_fan = wb.create_sheet('风扇信息')
    fan_headers = ['Sys_ip', 'Sys_name', 'Chassis/Slot', 'Fan_id', 'Fan_statu', 'Alarm_notes']
    write_header(ws_fan, fan_headers)
    fan_row = 2
    for dd in all_device_data:
        for f in dd['fan_list']:
            alarm = '--'
            if f['Fan_statu'].upper() not in ('NORMAL', 'FAN-LESS'):
                alarm = '风扇异常'
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'Chassis/Slot': f.get('Chassis/Slot', ''), 'Fan_id': f.get('Fan_id', ''),
                'Fan_statu': f.get('Fan_statu', ''), 'Alarm_notes': alarm,
            }
            write_data_rows(ws_fan, [row_data], fan_headers, fan_row)
            fan_row += 1
    auto_width(ws_fan)

    # --- Sheet 9: 聚合统计 ---
    ws_agg = wb.create_sheet('聚合统计')
    agg_headers = ['Sys_ip', 'Sys_name', 'AGG_Interface', 'AGG_Mode', 'Selected_port',
                   'Unselected_port', 'Individual_type', 'Share_type', 'Unselected_Ports', 'Alarm_notes']
    write_header(ws_agg, agg_headers)
    agg_row = 2
    for dd in all_device_data:
        for a in dd['agg_list']:
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'AGG_Interface': a.get('AGG_Interface', ''),
                'AGG_Mode': a.get('AGG_Mode', ''),
                'Selected_port': a.get('Selected_port', 0),
                'Unselected_port': a.get('Unselected_port', 0),
                'Individual_type': a.get('Individual_type', 0),
                'Share_type': a.get('Share_type', ''),
                'Unselected_Ports': a.get('Unselected_Ports', ''),
                'Alarm_notes': a.get('Alarm_notes', '--'),
            }
            write_data_rows(ws_agg, [row_data], agg_headers, agg_row)
            agg_row += 1
    auto_width(ws_agg)

    # --- Sheet 10: 光功率 ---
    ws_opt = wb.create_sheet('光功率')
    opt_headers = ['Sys_ip', 'Sys_name', 'Interface', 'Temp(°C)', 'Voltage(V)', 'Bias(mA)',
                   'RX power(dBm)', 'TX power(dBm)', 'RX_Alarm', 'TX_Alarm', 'Alarm_notes']
    write_header(ws_opt, opt_headers)
    opt_row = 2
    for dd in all_device_data:
        for t in dd['transceiver_list']:
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'Interface': t.get('Interface', ''), 'Temp(°C)': t.get('Temp(°C)', ''),
                'Voltage(V)': t.get('Voltage(V)', ''), 'Bias(mA)': t.get('Bias(mA)', ''),
                'RX power(dBm)': t.get('RX power(dBm)', ''),
                'TX power(dBm)': t.get('TX power(dBm)', ''),
                'RX_Alarm': t.get('RX_Alarm', '--'),
                'TX_Alarm': t.get('TX_Alarm', '--'),
                'Alarm_notes': t.get('Alarm_notes', '--'),
            }
            write_data_rows(ws_opt, [row_data], opt_headers, opt_row)
            opt_row += 1
    auto_width(ws_opt)

    # --- Sheet 11: BFD状态 ---
    ws_bfd = wb.create_sheet('BFD状态')
    bfd_headers = ['Sys_ip', 'Sys_name', 'BFD_Interface', 'State', 'Alarm_notes']
    write_header(ws_bfd, bfd_headers)
    bfd_row = 2
    for dd in all_device_data:
        for b in dd['bfd_list']:
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'BFD_Interface': b.get('BFD_Interface', ''),
                'State': b.get('State', ''), 'Alarm_notes': b.get('Alarm_notes', '--'),
            }
            write_data_rows(ws_bfd, [row_data], bfd_headers, bfd_row)
            bfd_row += 1
    auto_width(ws_bfd)

    # --- Sheet 12: MAC漂移（仅近90天内） ---
    ws_mac = wb.create_sheet('MAC漂移')
    mac_headers = ['Sys_ip', 'Sys_name', 'Move_Mac', 'Move_Vlan', 'Current_Port',
                   'Source_Port', 'Move_Time', 'Move_Count', 'Alarm_notes']
    write_header(ws_mac, mac_headers)
    mac_row = 2
    mac_cutoff = datetime.datetime(2026, 7, 16) - datetime.timedelta(days=90)
    for dd in all_device_data:
        for m in dd['mac_move_list']:
            move_time_str = m.get('Move_Time', '')
            try:
                move_time = datetime.datetime.strptime(move_time_str, '%Y-%m-%d %H:%M:%S')
                if move_time < mac_cutoff:
                    continue
            except ValueError:
                continue
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'Move_Mac': m.get('Move_Mac', ''), 'Move_Vlan': m.get('Move_Vlan', ''),
                'Current_Port': m.get('Current_Port', ''),
                'Source_Port': m.get('Source_Port', ''),
                'Move_Time': m.get('Move_Time', ''),
                'Move_Count': m.get('Move_Count', 0),
                'Alarm_notes': m.get('Alarm_notes', '--'),
            }
            write_data_rows(ws_mac, [row_data], mac_headers, mac_row)
            mac_row += 1
    auto_width(ws_mac)

    # --- Sheet 13: 版本信息 ---
    ws_ver = wb.create_sheet('版本信息')
    ver_headers = ['Sys_ip', 'Sys_name', '系统版本', '软件版本', '设备型号']
    write_header(ws_ver, ver_headers)
    ver_row = 2
    for dd in all_device_data:
        row_data = {
            'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
            '系统版本': dd.get('系统版本', 'V7'), '软件版本': dd.get('软件版本', ''),
            '设备型号': dd.get('设备型号', ''),
        }
        write_data_rows(ws_ver, [row_data], ver_headers, ver_row)
        ver_row += 1
    auto_width(ws_ver)

    # --- Sheet 14: 单板条码 ---
    ws_bc = wb.create_sheet('单板条码')
    bc_headers = ['Sys_ip', 'Sys_name', 'Slot', 'Device_Name', 'Info']
    write_header(ws_bc, bc_headers)
    bc_row = 2
    for dd in all_device_data:
        for bc in dd['barcode_list']:
            row_data = {
                'Sys_ip': dd['Sys_ip'], 'Sys_name': dd['Sys_name'],
                'Slot': bc.get('Slot', ''), 'Device_Name': bc.get('Device_Name', ''),
                'Info': bc.get('Info', ''),
            }
            write_data_rows(ws_bc, [row_data], bc_headers, bc_row)
            bc_row += 1
    auto_width(ws_bc)

    wb.save(output_path)
    print(f"Excel 已生成: {output_path}")


# ============================================================================
# 第四部分：Word 巡检报告生成
# ============================================================================

def build_abnormal_data(all_device_data):
    result = {}
    inspection_device_info = []
    for dd in all_device_data:
        inspection_device_info.append([
            dd['Sys_name'], dd.get('系统版本', 'V7'),
            dd.get('软件版本', ''), dd.get('设备型号', '')
        ])
    # ── 智能异常检测 ──
    anomaly_summary = []
    for dd in all_device_data:
        raw = dd.get('logbuffer_raw', '')
        if raw:
            anomalies = detect_smart_anomalies(raw)
            for a in anomalies:
                anomaly_summary.append([
                    dd.get('设备型号', ''), dd['Sys_ip'], dd['Sys_name'],
                    a['type'], a['interface'], a['summary']
                ])
    result['智能异常检测'] = anomaly_summary

    summary_df = pd.DataFrame([build_summary_row(dd) for dd in all_device_data])

    cpu_abnormal = []
    for _, row in summary_df.iterrows():
        if '异常' in str(row.get('CPU占用', '')):
            cpu_abnormal.append([row['设备型号'], row['Sys_ip'], row['Sys_name'], row['CPU利用率']])
    result['CPU占用'] = cpu_abnormal

    mem_abnormal = []
    for _, row in summary_df.iterrows():
        if '异常' in str(row.get('内存占用', '')):
            mem_abnormal.append([row['设备型号'], row['Sys_ip'], row['Sys_name'], row['内存利用率']])
    result['内存占用'] = mem_abnormal

    board_abnormal = []
    for dd in all_device_data:
        for b in dd['board_list']:
            state = b['State'].upper()
            btype = b['Type'].upper()
            if state not in ('NORMAL', 'MASTER', 'SLAVE', 'STANDBY', 'ABSENT') and btype != 'NONE':
                summary = build_summary_row(dd)
                board_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                       b['Slot'], b['State']])
    result['设备/板卡状态'] = board_abnormal

    power_abnormal = []
    for dd in all_device_data:
        summary = build_summary_row(dd)
        if '异常' in summary.get('电源状态', ''):
            pwr_details = []
            for p in dd['power_list']:
                if p['State'].upper() != 'NORMAL':
                    chassis = p.get('Chassis/Slot', '')
                    pwr_state = p['State'].upper()
                    if pwr_state == 'FAULT':
                        pwr_details.append(f"电源位 {p['PowerID']}：Fault（电源故障，请检查设备电源）")
                    elif pwr_state == 'ABSENT':
                        pwr_details.append(f"电源位 {p['PowerID']}：Absent（未安装或未到位）")
                    else:
                        pwr_details.append(f"电源位 {p['PowerID']}：{p['State']}")
            if pwr_details:
                power_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                       '\n'.join(pwr_details)])
    result['电源状态'] = power_abnormal

    fan_abnormal = []
    for dd in all_device_data:
        summary = build_summary_row(dd)
        if '异常' in summary.get('风扇状态', ''):
            for f in dd['fan_list']:
                if f['Fan_statu'].upper() != 'NORMAL':
                    fan_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                         f.get('Chassis/Slot', ''), f['Fan_id'], f['Fan_statu']])
    result['风扇状态'] = fan_abnormal

    agg_abnormal = []
    for dd in all_device_data:
        summary = build_summary_row(dd)
        if '异常' in summary.get('聚合选中状态', ''):
            agg_details = []
            for a in dd['agg_list']:
                if a['Alarm_notes'] != '--':
                    agg_details.append(f"{a['AGG_Interface']}：{a['Alarm_notes']}")
            if agg_details:
                agg_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                     '\n'.join(agg_details)])
    result['聚合选中状态'] = agg_abnormal

    optical_abnormal = []
    for dd in all_device_data:
        summary = build_summary_row(dd)
        if '异常' in summary.get('光功率状态', ''):
            details = []
            for t in dd['transceiver_list']:
                if t['Alarm_notes'] not in ('--', '需手动查看'):
                    # alarm_notes 可能为 '收光低于或接近临界; 发光高于或接近临界' 组合
                    for alarm_note in t['Alarm_notes'].split('; '):
                        alarm_note = alarm_note.strip()
                        if alarm_note.startswith('收光'):
                            rest = alarm_note[2:]
                            details.append(f"{t['Interface']}：收光功率 {t['RX_Alarm']}，{rest}阈值")
                        elif alarm_note.startswith('发光'):
                            rest = alarm_note[2:]
                            details.append(f"{t['Interface']}：发光功率 {t['TX_Alarm']}，{rest}阈值")
                        else:
                            details.append(f"{t['Interface']}：{t['RX_Alarm']}，{alarm_note}")
            if details:
                optical_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                         '\n'.join(details)])
    result['光功率状态'] = optical_abnormal

    in_usage_abnormal = []
    for dd in all_device_data:
        summary = build_summary_row(dd)
        if '异常' in summary.get('接口接收利用率', ''):
            details = []
            for r in dd['interface_rate_in']:
                try:
                    if float(r['Usage'].replace('%', '')) > 60:
                        usage_str = r['Usage']
                        if not usage_str.endswith('%'):
                            usage_str += '%'
                        details.append(f"{r['Interface']}（{usage_str}）")
                except ValueError:
                    continue
            if details:
                in_usage_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                          '；'.join(details)])
    result['接口接收利用率'] = in_usage_abnormal

    out_usage_abnormal = []
    for dd in all_device_data:
        summary = build_summary_row(dd)
        if '异常' in summary.get('接口发送利用率', ''):
            details = []
            for r in dd['interface_rate_out']:
                try:
                    if float(r['Usage'].replace('%', '')) > 60:
                        usage_str = r['Usage']
                        if not usage_str.endswith('%'):
                            usage_str += '%'
                        details.append(f"{r['Interface']}（{usage_str}）")
                except ValueError:
                    continue
            if details:
                out_usage_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                           '；'.join(details)])
    result['接口发送利用率'] = out_usage_abnormal

    in_err_abnormal = []
    for dd in all_device_data:
        for c in dd['interface_counters_in']:
            if c['Err_pkts'] > 0:
                summary = build_summary_row(dd)
                in_err_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                        f"接口:{c['Interface']},\n接收错包数:{c['Err_pkts']}"])
    result['接口接收错包'] = in_err_abnormal

    out_err_abnormal = []
    for dd in all_device_data:
        for c in dd['interface_counters_out']:
            if c['Err_pkts'] > 0:
                summary = build_summary_row(dd)
                out_err_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                         f"接口:{c['Interface']},\n发送错包数:{c['Err_pkts']}"])
    result['接口发送错包'] = out_err_abnormal

    bfd_abnormal = []
    for dd in all_device_data:
        for b in dd['bfd_list']:
            if 'down' in b.get('Alarm_notes', '').lower():
                summary = build_summary_row(dd)
                bfd_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                     b.get('BFD_Interface', ''), b.get('Alarm_notes', '')])
    result['BFD状态'] = bfd_abnormal

    mac_abnormal = []
    cutoff_date = datetime.datetime.now() - datetime.timedelta(days=90)
    for dd in all_device_data:
        if dd['mac_move_list']:
            summary = build_summary_row(dd)
            recent_moves = []
            for m in dd['mac_move_list']:
                move_time_str = m.get('Move_Time', '')
                try:
                    move_time = datetime.datetime.strptime(move_time_str, '%Y-%m-%d %H:%M:%S')
                    if move_time >= cutoff_date:
                        recent_moves.append(m)
                except ValueError:
                    continue
            if not recent_moves:
                continue
            recent_moves_sorted = sorted(recent_moves,
                                         key=lambda m: int(m.get('Move_Count', 0)),
                                         reverse=True)
            top_moves = recent_moves_sorted[:5]
            mac_str_list = []
            for m in top_moves:
                move_time = m.get('Move_Time', '')
                mac_str_list.append(
                    f"MAC:{m['Move_Mac']},VLAN:{m['Move_Vlan']},时间:{move_time},漂移次数:{m['Move_Count']}"
                )
            mac_abnormal.append([summary['设备型号'], summary['Sys_ip'], summary['Sys_name'],
                                 '；'.join(mac_str_list)])
    result['MAC漂移'] = mac_abnormal

    result['ERROR及以上日志'] = []
    result['巡检设备信息'] = inspection_device_info
    return result


def add_custom_heading_styles(doc):
    style = doc.styles.add_style('Heading1Custom', 1)
    style.base_style = doc.styles['Heading 1']
    style.font.name = '黑体'
    style.font.size = Pt(24)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '黑体')
    pPr = style.element.get_or_add_pPr()
    outline_lvl = OxmlElement('w:outlineLvl')
    outline_lvl.set(qn('w:val'), '0')
    pPr.append(outline_lvl)

    style = doc.styles.add_style('Heading2Custom', 1)
    style.base_style = doc.styles['Heading 2']
    style.font.name = '宋体'
    style.font.size = Pt(20)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '宋体')
    pPr = style.element.get_or_add_pPr()
    outline_lvl = OxmlElement('w:outlineLvl')
    outline_lvl.set(qn('w:val'), '1')
    pPr.append(outline_lvl)

    style = doc.styles.add_style('Heading3Custom', 1)
    style.base_style = doc.styles['Heading 3']
    style.font.name = '楷体'
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '楷体')
    pPr = style.element.get_or_add_pPr()
    outline_lvl = OxmlElement('w:outlineLvl')
    outline_lvl.set(qn('w:val'), '2')
    pPr.append(outline_lvl)

    style = doc.styles.add_style('toc1Custom', 1)
    style.font.name = '黑体'
    style.font.size = Pt(24)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '黑体')
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    style = doc.styles.add_style('toc2Custom', 1)
    style.font.name = '宋体'
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(4)

    style = doc.styles.add_style('toc3Custom', 1)
    style.font.name = '楷体'
    style.font.size = Pt(14)
    style.font.bold = False
    style.font.color.rgb = RGBColor(0, 0, 0)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '楷体')
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(2)


def create_toc(doc):
    paragraph = doc.add_paragraph()
    run1 = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fldChar_begin)
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \t "Heading1Custom,1,Heading2Custom,2,Heading3Custom,3" \h \z'
    run2._r.append(instrText)
    run3 = paragraph.add_run()
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar_separate)
    run4 = paragraph.add_run('（右键点击此处选择\u201c更新域\u201d生成目录）')
    run5 = paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar_end)


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_warning_header(doc, text, level='重要', color=RGBColor(255, 0, 0)):
    p = doc.add_paragraph('')
    run = p.add_run(f'级别：{level}')
    run.font.name = '仿宋'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = color
    p = doc.add_paragraph(f'\n{text}\n')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_table_with_data(doc, data, headers, table_title_color='ADD8E6', table_rows_color='E0F0F5'):
    if not data:
        p = doc.add_paragraph('\n无异常设备。\n')
        return
    table = doc.add_table(rows=len(data) + 1, cols=len(headers), style='Table Grid')
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        set_cell_shading(table.rows[0].cells[j], table_title_color)
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            if j < len(headers):
                if isinstance(val, float):
                    table.rows[i + 1].cells[j].text = str(int(val))
                else:
                    table.rows[i + 1].cells[j].text = str(val)
        if (i + 1) % 2 == 0:
            for j in range(len(headers)):
                set_cell_shading(table.rows[i + 1].cells[j], table_rows_color)


def generate_word_report(error_info, output_path, cover_image_path=None):
    doc = Document()
    add_custom_heading_styles(doc)

    section1 = doc.sections[0]
    section1.page_width = Inches(8.48)
    section1.page_height = Inches(11)
    section1.left_margin = Inches(0)
    section1.top_margin = Inches(0)
    section1.right_margin = Inches(0)
    section1.bottom_margin = Inches(0)

    # Delete any default empty paragraphs before adding the cover
    for p in list(doc.paragraphs):
        if not p.text.strip():
            p._element.getparent().remove(p._element)

    report_date = datetime.datetime.now()
    date_str = f"{report_date.year}年{report_date.month:02d}月{report_date.day:02d}日"

    cover_img_path = cover_image_path if cover_image_path and os.path.exists(cover_image_path) else None

    if cover_img_path:
        # --- Cover background image (floating, behind text) ---
        cover_img = doc.add_paragraph()
        cover_img.paragraph_format.space_before = Pt(0)
        cover_img.paragraph_format.space_after = Pt(0)
        cover_img.paragraph_format.line_spacing = Pt(1)
        run = cover_img.add_run()
        run.add_picture(cover_img_path, width=Inches(8.48), height=Inches(11))

        # Convert wp:inline → wp:anchor (floating behind text, zero flow space)
        drawing = run._element.find(qn('w:drawing'))
        inline = drawing.find(qn('wp:inline'))
        extent = inline.find(qn('wp:extent'))
        docPr_elem = inline.find(qn('wp:docPr'))
        cNvGraphicFramePr = inline.find(qn('wp:cNvGraphicFramePr'))
        graphic = inline.find(qn('a:graphic'))

        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '0')
        anchor.set('distR', '0')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '0')
        anchor.set('behindDoc', '1')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')

        simplePos = OxmlElement('wp:simplePos')
        simplePos.set('x', '0')
        simplePos.set('y', '0')
        anchor.append(simplePos)

        positionH = OxmlElement('wp:positionH')
        positionH.set('relativeFrom', 'page')
        ph_off = OxmlElement('wp:posOffset')
        ph_off.text = '0'
        positionH.append(ph_off)
        anchor.append(positionH)

        positionV = OxmlElement('wp:positionV')
        positionV.set('relativeFrom', 'page')
        pv_off = OxmlElement('wp:posOffset')
        pv_off.text = '0'
        positionV.append(pv_off)
        anchor.append(positionV)

        anchor.append(extent)

        effectExtent = OxmlElement('wp:effectExtent')
        effectExtent.set('l', '0')
        effectExtent.set('t', '0')
        effectExtent.set('r', '0')
        effectExtent.set('b', '0')
        anchor.append(effectExtent)

        wrapNone = OxmlElement('wp:wrapNone')
        anchor.append(wrapNone)

        anchor.append(docPr_elem)
        anchor.append(cNvGraphicFramePr)
        anchor.append(graphic)

        drawing.remove(inline)
        drawing.append(anchor)

    def add_cover_textbox(doc, text_lines, y_twips, w_twips=7200, h_twips=2880, dark_bg=True):
        """Add a single w:framePr floating paragraph containing all cover text lines."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(30)
        pPr = p._element.get_or_add_pPr()
        framePr = OxmlElement('w:framePr')
        framePr.set(qn('w:w'), str(w_twips))
        framePr.set(qn('w:h'), str(h_twips))
        framePr.set(qn('w:hRule'), 'atLeast')
        framePr.set(qn('w:xAlign'), 'center')
        framePr.set(qn('w:y'), str(y_twips))
        framePr.set(qn('w:wrap'), 'none')
        framePr.set(qn('w:hAnchor'), 'page')
        framePr.set(qn('w:vAnchor'), 'page')
        pPr.append(framePr)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        color = RGBColor(255, 255, 255) if dark_bg else RGBColor(0, 0, 0)
        for i, (text, font_size, bold) in enumerate(text_lines):
            if i > 0:
                br_run = p.add_run()
                br_run.add_break()
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.color.rgb = color
        return p

    # 无论是否有封面图，都添加标题文字层
    add_cover_textbox(
        doc,
        [
            ('H3C 网络设备巡检总结报告', 26, True),
            ('AIOps 智能运维托管平台', 18, False),
            (date_str, 14, False),
        ],
        y_twips=11520,
        dark_bg=bool(cover_img_path),
    )

    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    section2.page_width = Inches(8.5)
    section2.page_height = Inches(11)
    section2.left_margin = Inches(1)
    section2.right_margin = Inches(1)
    section2.top_margin = Inches(1)
    section2.bottom_margin = Inches(1)

    toc_text = doc.add_paragraph('目录')
    toc_text.style = 'toc1Custom'
    create_toc(doc)
    doc.add_page_break()

    table_title_color = 'ADD8E6'
    table_rows_color = 'E0F0F5'

    doc.add_paragraph('一、设备健康状态巡查', style='Heading1Custom')
    doc.add_paragraph('1.1、硬件状态巡查', style='Heading2Custom')

    doc.add_paragraph('1.1.1、单板状态巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备单板均为normal状态，无fault信息', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    board_data = error_info.get('设备/板卡状态', [])
    if board_data:
        add_table_with_data(doc, board_data, ['设备型号', '设备IP', '设备名称', 'Slot', 'Brd_statu'])
    else:
        doc.add_paragraph('无异常设备。')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('建议逐项检查单板是否存在重启记录、主控版本一致性与供电状态；若以上均正常，请联系 H3C 厂商技术支持进一步诊断。\n')
    doc.add_page_break()

    doc.add_paragraph('1.1.2、电源状态巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备电源均在位，且无fault状态', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    power_data = error_info.get('电源状态', [])
    if power_data:
        add_table_with_data(doc, power_data, ['设备型号', '设备IP', '设备名称', '电源异常信息'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('请检查power状态：Absent状态为未安装电源，或电源安装不到位，Fault状态为电源故障或电源未供电。\n')
    doc.add_page_break()

    doc.add_paragraph('1.1.3、风扇状态巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备风扇均在位，且无fault状态', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    fan_data = error_info.get('风扇状态', [])
    if fan_data:
        add_table_with_data(doc, fan_data, ['设备型号', '设备IP', '设备名称', 'Chassis/Slot', 'fan id', 'fan 状态'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('建议检查风扇物理安装是否到位、接口是否松动、供电是否正常；若以上均正常，请联系 H3C 厂商技术支持进一步诊断。\n')
    doc.add_page_break()

    doc.add_paragraph('1.2、软件状态巡查', style='Heading2Custom')

    doc.add_paragraph('1.2.1、CPU状态巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备CPU利用率不得高于70%', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    cpu_data = error_info.get('CPU占用', [])
    if cpu_data:
        add_table_with_data(doc, cpu_data, ['设备型号', '设备IP', '设备名称', 'CPU当前利用率'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('请检查CPU占用率高的任务，是否有报文攻击、环路、路由振荡等异常情况。\n')
    doc.add_page_break()

    doc.add_paragraph('1.2.2、内存状态巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备内存利用率不得高于80%', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    mem_data = error_info.get('内存占用', [])
    if mem_data:
        add_table_with_data(doc, mem_data, ['设备型号', '设备IP', '设备名称', '内存当前利用率'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('请检查mem占用率高的任务，是否有报文攻击、环路、路由振荡等异常情况。\n')
    doc.add_page_break()

    doc.add_paragraph('1.2.3、ERROR级别以上日志', style='Heading3Custom')
    add_warning_header(doc, '要求设备无ERROR级别以上日志', '一般', RGBColor(0, 0, 255))
    doc.add_paragraph('异常设备信息：')
    anomaly_data = error_info.get('智能异常检测', [])
    if anomaly_data:
        add_table_with_data(doc, anomaly_data, ['设备型号', '设备IP', '设备名称', '异常类型', '接口', '异常摘要'])
    else:
        doc.add_paragraph('\n未发现异常日志。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('日志条目过多时，请登录设备进一步检查。\n')
    doc.add_page_break()

    doc.add_paragraph('1.3、协议状态巡查', style='Heading2Custom')

    doc.add_paragraph('1.3.1、聚合组状态巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备上聚合组状态均正常，链路均选中', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    agg_data = error_info.get('聚合选中状态', [])
    if agg_data:
        add_table_with_data(doc, agg_data, ['设备型号', '设备IP', '设备名称', '聚合组编号及异常信息'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('建议检查聚合组配置：确认物理端口状态、两端 LACP 协商模式是否一致，排查是否存在遗留配置；若上述检查均无异常，请联系网络技术负责人协查。\n')
    doc.add_page_break()

    doc.add_paragraph('1.3.2、BFD状态巡查', style='Heading3Custom')
    add_warning_header(doc, '要求已启用BFD的设备BFD状态均正常', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    bfd_data = error_info.get('BFD状态', [])
    if bfd_data:
        add_table_with_data(doc, bfd_data, ['设备型号', '设备IP', '设备名称', '对应接口', '异常信息'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('建议检查异常 BFD 会话的关联接口状态、探测可达性及路由表项；确认对端 BFD 配置是否正常；若上述检查均无异常，请联系网络技术负责人协查。\n')
    doc.add_page_break()

    doc.add_paragraph('1.3.3、MAC地址漂移巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备上无漂移信息', '一般', RGBColor(0, 0, 255))
    doc.add_paragraph('异常设备信息：')
    mac_data = error_info.get('MAC漂移', [])
    if mac_data:
        add_table_with_data(doc, mac_data, ['设备型号', '设备IP', '设备名称', 'MAC漂移统计信息'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('漂移条目过多时，请登录设备进一步检查。\n')
    doc.add_page_break()

    doc.add_paragraph('1.4、接口状态巡查', style='Heading2Custom')

    doc.add_paragraph('1.4.1、光功率巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备上光模块功率均在合理范围内', '一般', RGBColor(0, 0, 255))
    doc.add_paragraph('异常设备信息：')
    opt_data = error_info.get('光功率状态', [])
    if opt_data:
        add_table_with_data(doc, opt_data, ['设备型号', '设备IP', '设备名称', '光功率异常信息'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('请检查两端收发光功率，确认线缆质量、实际光纤距离等因素；若光模块及光纤链路均无异常，请联系 H3C 厂商技术支持进一步诊断。\n')
    doc.add_page_break()

    doc.add_paragraph('1.4.2、接口出方向利用率巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备上接口出方向流量低于60%', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    out_usage_data = error_info.get('接口发送利用率', [])
    if out_usage_data:
        add_table_with_data(doc, out_usage_data, ['设备型号', '设备IP', '设备名称', '出方向利用率'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('建议出/入方向利用率控制在 60% 以下；双链路聚合场景下，单链路带宽占用建议不超过 50%。\n')
    doc.add_page_break()

    doc.add_paragraph('1.4.3、接口入方向利用率巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备上接口入方向流量低于60%', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    in_usage_data = error_info.get('接口接收利用率', [])
    if in_usage_data:
        add_table_with_data(doc, in_usage_data, ['设备型号', '设备IP', '设备名称', '入方向利用率'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    doc.add_paragraph('建议出/入方向利用率控制在 60% 以下；双链路聚合场景下，单链路带宽占用建议不超过 50%。\n')
    doc.add_page_break()

    doc.add_paragraph('1.4.4、接口出方向错包增长巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备上采集期间出方向无错包', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    out_err_data = error_info.get('接口发送错包', [])
    if out_err_data:
        add_table_with_data(doc, out_err_data, ['设备型号', '设备IP', '设备名称', '接口出方向错包'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    if out_err_data:
        doc.add_paragraph('已发现出方向错包增长，建议首先确认错包类型（CRC、Overrun、Jabber 等），针对性排查物理链路或端口协商问题。若为 CRC 错包增长，检查光模块、光纤链路及接口双工模式；若为 Overrun 错包，检查接口缓冲区配置及突发流量。\n')
    else:
        doc.add_paragraph('本次采集期间未发现出方向错包增长，无需处置。如后续出现错包，建议首先确认错包类型（CRC、Overrun、Jabber 等），针对性排查物理链路或端口协商问题。\n')
    doc.add_page_break()

    doc.add_paragraph('1.4.5、接口入方向错包增长巡查', style='Heading3Custom')
    add_warning_header(doc, '要求设备上采集期间入方向无错包', '重要', RGBColor(255, 0, 0))
    doc.add_paragraph('异常设备信息：')
    in_err_data = error_info.get('接口接收错包', [])
    if in_err_data:
        add_table_with_data(doc, in_err_data, ['设备型号', '设备IP', '设备名称', '入方向错包'])
    else:
        doc.add_paragraph('\n无异常设备。\n')
    doc.add_paragraph('\n改进建议：')
    if in_err_data:
        doc.add_paragraph('已发现入方向错包增长，建议首先确认错包类型（CRC、Overrun、Jabber 等），针对性排查物理链路或端口协商问题。若为 CRC 错包增长，检查光模块、光纤链路及接口双工模式；若为 Overrun 错包，检查接口缓冲区配置及突发流量。\n')
    else:
        doc.add_paragraph('本次采集期间未发现入方向错包增长，无需处置。如后续出现错包，建议首先确认错包类型（CRC、Overrun、Jabber 等），针对性排查物理链路或端口协商问题。\n')
    doc.add_page_break()

    doc.add_paragraph('二、附录文件', style='Heading1Custom')
    doc.add_paragraph('2.1、设备清单', style='Heading2Custom')
    device_info_data = error_info.get('巡检设备信息', [])
    if device_info_data:
        add_table_with_data(doc, device_info_data, ['系统名', '系统版本', '软件版本', '设备型号'])

    doc.save(output_path)
    print(f"Word 报告已生成: {output_path}")


# ============================================================================
# 第五部分：主入口（多进程并行解析）
# ============================================================================

def auto_update_toc(docx_path):
    """
    平台环境下不依赖 win32com，仅保留占位日志。
    Word 打开后会提示用户右键目录 → 更新域以显示页码。
    """
    # Linux 服务器通常没有 Word/win32com，保持域代码即可
    pass


def generate_reports(all_device_data, output_dir, prefix="", cover_image_path=None):
    """供后端服务调用的统一报告生成入口。

    Args:
        all_device_data: list[dict] 解析后的设备数据
        output_dir: 报告输出目录
        prefix: 文件名前缀
        cover_image_path: Word 封面背景图路径，不存在则使用文字封面

    Returns:
        (xlsx_path, docx_path)
    """
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d-%H-%M-%S')
    filename_prefix = f"{prefix}-{timestamp}" if prefix else timestamp

    xlsx_path = os.path.join(output_dir, f"{filename_prefix}.xlsx")
    generate_excel(all_device_data, xlsx_path)

    error_info = build_abnormal_data(all_device_data)
    docx_path = os.path.join(output_dir, f"{filename_prefix}.docx")
    generate_word_report(error_info, docx_path, cover_image_path=cover_image_path)

    return xlsx_path, docx_path
